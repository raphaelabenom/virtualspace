import ast
import astroid
from jinja2 import Environment, FileSystemLoader
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE

class AutoDoc:
    def __init__(self, source_dir):
        self.source_dir = source_dir
        self.ast_trees = {}
        self.astroid_trees = {}
        self.classes = []
        self.functions = []
        self.modules = []

    def analyze_code(self):
        """Analisa todos os arquivos Python no diretório especificado."""
        for root, _, files in os.walk(self.source_dir):
            for file in files:
                if file.endswith('.py'):
                    file_path = os.path.join(root, file)
                    with open(file_path, 'r') as source_file:
                        source_code = source_file.read()
                        self.ast_trees[file_path] = ast.parse(source_code)
                        self.astroid_trees[file_path] = astroid.parse(source_code)

    def extract_information(self):
        """Extrai informações de módulos, classes e funções dos arquivos analisados."""
        for file_path, tree in self.astroid_trees.items():
            module = self._extract_module_info(tree, file_path)
            self.modules.append(module)
            self.classes.extend(self._extract_classes(tree))
            self.functions.extend(self._extract_functions(tree))

    def _extract_module_info(self, tree, file_path):
        """Extrai informações do módulo."""
        return {
            'name': os.path.basename(file_path),
            'path': file_path,
            'docstring': ast.get_docstring(tree),
        }

    def _extract_classes(self, tree):
        """Extrai informações das classes no módulo."""
        classes = []
        for node in tree.body:
            if isinstance(node, astroid.ClassDef):
                classes.append({
                    'name': node.name,
                    'docstring': node.doc_node.value if node.doc_node else '',
                    'methods': self._extract_methods(node),
                    'attributes': self._extract_attributes(node),
                })
        return classes

    def _extract_methods(self, class_node):
        """Extrai informações dos métodos de uma classe."""
        methods = []
        for node in class_node.body:
            if isinstance(node, astroid.FunctionDef):
                methods.append({
                    'name': node.name,
                    'docstring': node.doc_node.value if node.doc_node else '',
                    'parameters': [param.name for param in node.args.args],
                })
        return methods

    def _extract_attributes(self, class_node):
        """Extrai atributos de uma classe."""
        attributes = []
        for node in class_node.body:
            if isinstance(node, astroid.Assign):
                for target in node.targets:
                    if isinstance(target, astroid.AssignName):
                        attributes.append(target.name)
        return attributes

    def _extract_functions(self, tree):
        """Extrai informações das funções no módulo."""
        functions = []
        for node in tree.body:
            if isinstance(node, astroid.FunctionDef):
                functions.append({
                    'name': node.name,
                    'docstring': node.doc_node.value if node.doc_node else '',
                    'parameters': [param.name for param in node.args.args],
                })
        return functions

    def generate_documentation(self):
        """Gera a documentação em formato Markdown."""
        env = Environment(loader=FileSystemLoader('templates'))
        template = env.get_template('documentation_template.md')
        
        context = {
            'modules': self.modules,
            'classes': self.classes,
            'functions': self.functions,
        }
        
        documentation = template.render(context)
        
        with open('documentation.md', 'w') as doc_file:
            doc_file.write(documentation)

    def generate_uml_diagrams(self):
        """Gera diagramas UML usando Mermaid."""
        class_diagram = """
        classDiagram
        {% for class in self.classes %}
        class {{ class.name }} {
            {% for attribute in class.attributes %}
            +{{ attribute }}
            {% endfor %}
            {% for method in class.methods %}
            +{{ method.name }}({% for param in method.parameters %}{{ param }}{% if not loop.last %}, {% endif %}{% endfor %})
            {% endfor %}
        }
        {% endfor %}
        """
        
        # Adicionar relações entre classes
        for class_ in self.classes:
            for method in class_['methods']:
                for param in method['parameters']:
                    if param in [c['name'] for c in self.classes]:
                        class_diagram += f"{class_['name']} --> {param}\n"
        
        with open('class_diagram.mmd', 'w') as f:
            f.write(class_diagram)

    def generate_mermaid_graphs(self):
        """Gera gráficos de dependência usando Mermaid."""
        dependency_graph = """
        graph TD
        {% for module in self.modules %}
        {{ module.name }}
        {% endfor %}
        {% for class in self.classes %}
        {{ class.name }}
        {% endfor %}
        """
        
        # Adicionar relações de dependência
        for class_ in self.classes:
            for method in class_['methods']:
                for param in method['parameters']:
                    if param in [c['name'] for c in self.classes]:
                        dependency_graph += f"{class_['name']} --> {param}\n"
        
        # Adicionar relações entre módulos e classes
        for module in self.modules:
            for class_ in self.classes:
                if class_['name'] in module['path']:
                    dependency_graph += f"{module['name']} --> {class_['name']}\n"
        
        with open('dependency_graph.mmd', 'w') as f:
            f.write(dependency_graph)

    def extract_requirements(self):
        """Extrai requisitos dos docstrings."""
        requirements = []
        for module in self.modules:
            if module['docstring']:
                requirements.extend(self._extract_requirements_from_docstring(module['docstring']))
        for class_ in self.classes:
            if class_['docstring']:
                requirements.extend(self._extract_requirements_from_docstring(class_['docstring']))
        for function in self.functions:
            if function['docstring']:
                requirements.extend(self._extract_requirements_from_docstring(function['docstring']))
        
        with open('requirements.md', 'w') as f:
            f.write("# Requisitos do Sistema\n\n")
            for i, req in enumerate(requirements, 1):
                f.write(f"{i}. {req}\n")

    def _extract_requirements_from_docstring(self, docstring):
        """Extrai requisitos de um docstring."""
        requirements = []
        keywords = ['deve', 'precisa', 'requer', 'necessita', 'tem que', 'é necessário']
        sentences = re.split(r'[.!?]+', docstring)
        for sentence in sentences:
            if any(keyword in sentence.lower() for keyword in keywords):
                requirements.append(sentence.strip())
        return requirements

    def export_to_docx(self, output_file='documentation.docx'):
        """Exporta a documentação para um arquivo DOCX."""
        doc = Document()

        # Configurar estilos
        styles = doc.styles
        title_style = styles['Title'] if 'Title' in styles else styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(24)
        title_style.font.bold = True

        heading1_style = styles['Heading 1'] if 'Heading 1' in styles else styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.size = Pt(18)
        heading1_style.font.bold = True

        heading2_style = styles['Heading 2'] if 'Heading 2' in styles else styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.font.size = Pt(16)
        heading2_style.font.bold = True

        # Título
        doc.add_paragraph('Documentação Automática', style='Title')

        # Resumo do Sistema
        doc.add_paragraph('Resumo do Sistema', style='Heading 1')
        doc.add_paragraph('Este documento foi gerado automaticamente para fornecer uma visão geral do sistema analisado.')

        # Módulos
        doc.add_paragraph('Módulos', style='Heading 1')
        for module in self.modules:
            doc.add_paragraph(module['name'], style='Heading 2')
            doc.add_paragraph(f"Caminho: {module['path']}")
            if module['docstring']:
                doc.add_paragraph(module['docstring'])

        # Classes
        doc.add_paragraph('Classes', style='Heading 1')
        for class_ in self.classes:
            doc.add_paragraph(class_['name'], style='Heading 2')
            if class_['docstring']:
                doc.add_paragraph(class_['docstring'])
            
            doc.add_paragraph('Atributos:', style='Heading 2')
            for attr in class_['attributes']:
                doc.add_paragraph(attr, style='List Bullet')
            
            doc.add_paragraph('Métodos:', style='Heading 2')
            for method in class_['methods']:
                doc.add_paragraph(method['name'], style='Heading 2')
                if method['docstring']:
                    doc.add_paragraph(method['docstring'])
                doc.add_paragraph('Parâmetros:')
                for param in method['parameters']:
                    doc.add_paragraph(param, style='List Bullet')

        # Funções
        doc.add_paragraph('Funções', style='Heading 1')
        for function in self.functions:
            doc.add_paragraph(function['name'], style='Heading 2')
            if function['docstring']:
                doc.add_paragraph(function['docstring'])
            doc.add_paragraph('Parâmetros:')
            for param in function['parameters']:
                doc.add_paragraph(param, style='List Bullet')

        # Diagramas UML
        doc.add_paragraph('Diagramas UML', style='Heading 1')
        doc.add_paragraph('(Os diagramas UML gerados podem ser encontrados nos arquivos .mmd separados)')

        # Gráficos de Relações
        doc.add_paragraph('Gráficos de Relações', style='Heading 1')
        doc.add_paragraph('(Os gráficos de relações gerados podem ser encontrados nos arquivos .mmd separados)')

        # Requisitos do Sistema
        doc.add_paragraph('Requisitos do Sistema', style='Heading 1')
        with open('requirements.md', 'r') as req_file:
            requirements = req_file.read()
        doc.add_paragraph(requirements)

        # Salvar o documento
        doc.save(output_file)
        print(f"Documentação exportada para {output_file}")

if __name__ == "__main__":
    autodoc = AutoDoc("../compare_models.py")
    autodoc.analyze_code()
    autodoc.extract_information()
    autodoc.generate_documentation()
    autodoc.generate_uml_diagrams()
    autodoc.generate_mermaid_graphs()
    autodoc.extract_requirements()
    autodoc.export_to_docx()
    print("Documentação gerada e exportada com sucesso!")

